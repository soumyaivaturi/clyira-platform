"""
Document Service — Handles document upload, extraction, classification, and AI creation.
"""
import os
import re
import logging
from typing import Optional

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.core.config import settings
from app.models.document import Document, DocumentReference
from app.dtap import DTAPRegistry

logger = logging.getLogger(__name__)


class DocumentService:
    """Service layer for document operations"""

    def __init__(self, db: AsyncSession):
        self.db = db

    async def upload_document(
        self,
        file_content: bytes,
        filename: str,
        company_id: str,
        user_id: str,
        metadata: dict,
    ) -> Document:
        """
        Process an uploaded document:
        1. Compute content hash for duplicate detection (§22.3)
        2. Extract text content
        3. Classify document type
        4. Assign DTAP
        5. Store in database with hash
        """
        import hashlib
        from fastapi import HTTPException

        file_type = filename.rsplit(".", 1)[-1].lower() if "." in filename else "unknown"

        # Compute SHA-256 of raw bytes for dedup check (§22.3)
        content_hash = hashlib.sha256(file_content).hexdigest()
        existing = await self.db.execute(
            select(Document)
            .where(Document.company_id == company_id)
            .where(Document.content_hash == content_hash)
            .limit(1)
        )
        existing_doc = existing.scalar_one_or_none()
        if existing_doc:
            raise HTTPException(
                status_code=409,
                detail={
                    "code": "duplicate_document",
                    "message": "This file has already been uploaded.",
                    "existing_document_id": existing_doc.id,
                    "existing_document_title": existing_doc.title,
                },
            )

        # Run IDP extraction (Phase 3) — provides structured page/table/field data
        idp_sections: dict = {}
        if file_type == "pdf":
            try:
                from app.services.idp_engine import IDPEngine
                idp_out = IDPEngine().extract_to_dict(file_content, filename)
                idp_sections = idp_out
            except Exception as idp_err:
                logger.warning(f"IDP extraction failed for {filename}: {idp_err}")

        # Extract text from bytes in memory (before any storage write)
        extracted_text = await self._extract_text_from_bytes(file_content, file_type)
        # Merge IDP sections with rule-based section identification
        sections_from_text = self._identify_sections(extracted_text)
        extracted_sections = {**sections_from_text, "_idp": idp_sections} if idp_sections else sections_from_text

        # Save file to storage (Supabase in prod, local in dev)
        file_path = self._save_file(file_content, filename, company_id)

        # Auto-classify if metadata not provided
        document_category = metadata.get("document_category", "")
        if not document_category:
            document_category = self._auto_classify(filename, extracted_text)

        # Resolve DTAP
        dtap = DTAPRegistry.get_by_category(document_category)
        dtap_id = dtap.dtap_id if dtap else None

        # Create record
        document = Document(
            company_id=company_id,
            uploaded_by=user_id,
            title=metadata.get("title", filename),
            document_number=metadata.get("document_number"),
            version=metadata.get("version", "1.0"),
            function_type=metadata.get("function_type"),
            regulatory_category=metadata.get("regulatory_category"),
            department_owner=metadata.get("department_owner"),
            document_category=document_category,
            regulatory_frameworks=metadata.get("regulatory_frameworks"),
            dtap_id=dtap_id,
            file_path=file_path,
            file_type=file_type,
            file_size_bytes=len(file_content),
            content_hash=content_hash,
            extracted_text=extracted_text,
            extracted_sections=extracted_sections,
            status="ready",
        )

        self.db.add(document)
        await self.db.commit()
        await self.db.refresh(document)
        return document

    async def add_reference(
        self,
        document_id: str,
        file_content: bytes,
        filename: str,
        user_id: str,
        title: str,
        description: str = "",
        reference_type: str = "organizational_guideline",
    ) -> DocumentReference:
        """
        Add a user-uploaded reference to a document.
        These references are used during assessment as organizational context.
        """
        # Get the parent document for company context
        doc = await self.db.get(Document, document_id)
        if not doc:
            raise ValueError(f"Document {document_id} not found")

        file_type = filename.rsplit(".", 1)[-1].lower() if "." in filename else "unknown"

        # Extract text from bytes in memory
        extracted_text = await self._extract_text_from_bytes(file_content, file_type)

        # Save file to storage (Supabase in prod, local in dev)
        file_path = self._save_file(file_content, filename, doc.company_id, subfolder="references")

        reference = DocumentReference(
            document_id=document_id,
            uploaded_by=user_id,
            title=title or filename,
            description=description,
            file_path=file_path,
            file_type=file_type,
            extracted_text=extracted_text,
            reference_type=reference_type,
        )

        self.db.add(reference)
        await self.db.commit()
        await self.db.refresh(reference)
        return reference

    async def get_document_with_references(self, document_id: str) -> dict:
        """Get document and its references for assessment"""
        doc = await self.db.get(Document, document_id)
        if not doc:
            return None

        # Load references
        result = await self.db.execute(
            select(DocumentReference).where(DocumentReference.document_id == document_id)
        )
        references = result.scalars().all()

        return {
            "document": doc,
            "references": [
                {
                    "id": ref.id,
                    "title": ref.title,
                    "extracted_text": ref.extracted_text,
                    "reference_type": ref.reference_type,
                }
                for ref in references
            ],
        }

    def _save_file(self, content: bytes, filename: str, company_id: str, subfolder: str = "documents") -> str:
        """Upload to Supabase Storage in production, local filesystem in development."""
        if settings.SUPABASE_URL and settings.SUPABASE_SERVICE_KEY:
            from supabase import create_client
            storage_path = f"{company_id}/{subfolder}/{filename}"
            client = create_client(settings.SUPABASE_URL, settings.SUPABASE_SERVICE_KEY)
            client.storage.from_(settings.SUPABASE_STORAGE_BUCKET).upload(
                storage_path, content, {"upsert": "true"}
            )
            return storage_path
        storage_dir = os.path.join(settings.STORAGE_PATH, company_id, subfolder)
        os.makedirs(storage_dir, exist_ok=True)
        file_path = os.path.join(storage_dir, filename)
        with open(file_path, "wb") as f:
            f.write(content)
        return file_path

    async def _extract_text_from_bytes(self, content: bytes, file_type: str) -> str:
        """Extract text from document bytes in memory (no disk I/O)."""
        import io
        if file_type == "docx":
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(io.BytesIO(content))
                parts = []
                # Paragraphs (body text, headings)
                for p in doc.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)
                # Tables — test methods and lab docs are almost entirely table-formatted
                for table in doc.tables:
                    for row in table.rows:
                        row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_texts:
                            parts.append(" | ".join(row_texts))
                return "\n".join(parts)
            except Exception as e:
                logger.error(f"DOCX extraction failed: {e}")
                return ""
        elif file_type == "pdf":
            # Primary: pdfplumber — better for structured/table-heavy PDFs
            try:
                import pdfplumber
                pages = []
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text() or ""
                        # Also extract tables as pipe-delimited text
                        for table in page.extract_tables():
                            rows = [" | ".join(str(c or "").strip() for c in row) for row in table if any(c for c in row)]
                            text += "\n" + "\n".join(rows)
                        if text.strip():
                            pages.append(text)
                result = "\n".join(pages)
                if result.strip():
                    return result
            except Exception as e:
                logger.warning(f"pdfplumber failed, falling back to PyPDF2: {e}")
            # Fallback: PyPDF2
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(io.BytesIO(content))
                return "\n".join(page.extract_text() or "" for page in reader.pages)
            except Exception as e:
                logger.error(f"PDF extraction failed entirely: {e}")
                return ""
        elif file_type in ("txt", "md"):
            return content.decode("utf-8", errors="ignore")
        return ""

    def _identify_sections(self, text: str) -> dict:
        """Identify document sections from extracted text"""
        sections = {}
        current_section = "preamble"
        current_content = []

        for line in text.split("\n"):
            stripped = line.strip()
            # Heuristic: lines that are short, capitalized/numbered are likely headers
            if (
                stripped
                and len(stripped) < 100
                and (stripped.isupper() or stripped[0].isdigit() or stripped.endswith(":"))
            ):
                if current_content:
                    sections[current_section] = "\n".join(current_content)
                current_section = stripped
                current_content = []
            else:
                current_content.append(line)

        if current_content:
            sections[current_section] = "\n".join(current_content)

        return sections

    def _auto_classify(self, filename: str, text: str) -> str:
        """Auto-classify document category from filename and content"""
        filename_lower = filename.lower()
        text_lower = text[:2000].lower()

        # MBR/BPR — check before SOP to avoid false positives from "batch" appearing in SOPs
        if (
            any(kw in filename_lower for kw in ("mbr", "bpr", "batch_record", "batch-record", "batch_production", "dhr"))
            or "master batch record" in text_lower
            or "batch production record" in text_lower
            or "device history record" in text_lower
            or bool(re.search(r'\bbpr\b', text_lower))
            or bool(re.search(r'\bmbr\b', text_lower))
            or bool(re.search(r'\bdhr\b', text_lower))
            or (
                "batch record" in text_lower
                and any(kw in text_lower for kw in (
                    "lot number", "theoretical yield", "bill of materials",
                    "batch disposition", "in-process control", "manufacturing steps",
                    "executed by", "line clearance",
                ))
            )
        ):
            return "MBR"
        elif "sop" in filename_lower or "standard operating" in text_lower:
            return "SOP"
        elif "capa" in filename_lower or "corrective and preventive" in text_lower:
            return "CAPA"
        elif "atm" in filename_lower or "analytical test method" in text_lower or "test method" in text_lower:
            return "ATM"
        elif (
            any(kw in filename_lower for kw in ("coa", "cof", "certificate_of_analysis", "qc_test", "qc_record", "lab_report"))
            or "certificate of analysis" in text_lower
            or "certificate of conformance" in text_lower
            or bool(re.search(r'\bcoa\b', text_lower))
            or (
                "test results" in text_lower
                and any(kw in text_lower for kw in (
                    "specification", "acceptance criteria", "analyst", "system suitability",
                    "assay", "related substances", "dissolution", "microbial limits",
                ))
            )
        ):
            return "QC_TEST"
        elif "deviation" in filename_lower or "deviation report" in text_lower:
            return "Deviation"
        elif (
            "lir" in filename_lower
            or "laboratory incident" in text_lower
            or "lab incident" in text_lower
            or "laboratory investigation report" in text_lower
            or "out-of-specification" in text_lower
            or "oos investigation" in text_lower
        ):
            return "LIR"
        elif (
            "validation" in filename_lower
            or "installation qualification" in text_lower
            or "operational qualification" in text_lower
            or "performance qualification" in text_lower
            or "process validation" in text_lower
            or "cleaning validation" in text_lower
            or "computer system validation" in text_lower
            or re.search(r'\b(?:iq|oq|pq)\s*(?:protocol|report|validation)\b', text_lower)
        ):
            return "Validation"
        else:
            return "Other"
