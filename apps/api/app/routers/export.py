"""
Export Router — Generate downloadable DOCX assessment reports.
"""
import io
from datetime import datetime
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.core.database import get_db
from app.core.dependencies import get_current_user
from app.models.assessment import Assessment, Finding
from app.models.document import Document
from app.models.user import User

router = APIRouter()

# Severity ordering for report
_SEV_ORDER = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}
_LEVEL_NAMES = {
    "L1": "Structural Integrity", "L2": "Document Control",
    "L3": "Content Quality", "L4": "ALCOA+ Data Integrity",
    "L5": "Data Intelligence", "L6": "Cross-Document Consistency",
    "L7": "Lifecycle Compliance", "L8": "Regulatory Gap Analysis",
    "L9": "Enforcement Risk", "L10": "Longitudinal Intelligence",
    "L11": "Submission Readiness",
}


def _score_color(score: Optional[float]) -> str:
    if score is None:
        return "808080"
    if score >= 90:
        return "16A34A"
    if score >= 80:
        return "65A30D"
    if score >= 65:
        return "D97706"
    if score >= 50:
        return "DC2626"
    return "7F1D1D"


def _sev_color(sev: str) -> str:
    return {"critical": "DC2626", "high": "EA580C", "medium": "D97706", "low": "2563EB", "info": "6B7280"}.get(sev, "374151")


def _build_docx(
    assessment: Assessment,
    document: Document,
    findings: list[Finding],
) -> bytes:
    from docx import Document as DocxDoc
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = DocxDoc()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Helper: add colored heading ───────────────────────────────────────────
    def add_heading(text: str, level: int = 1, color: str = "1E3A5F"):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = RGBColor(
                int(color[:2], 16), int(color[2:4], 16), int(color[4:], 16)
            )
        return h

    def add_kv(para, key: str, value: str):
        run_k = para.add_run(f"{key}: ")
        run_k.bold = True
        run_k.font.size = Pt(10)
        run_v = para.add_run(value)
        run_v.font.size = Pt(10)

    def hex_to_rgb(h):
        return int(h[:2], 16), int(h[2:4], 16), int(h[4:], 16)

    # ── Cover ─────────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Clyira Assessment Report")
    title_run.bold = True
    title_run.font.size = Pt(22)
    r, g, b = hex_to_rgb("1E3A5F")
    title_run.font.color.rgb = RGBColor(r, g, b)

    doc.add_paragraph()  # spacer

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_name = meta.add_run(document.title)
    doc_name.bold = True
    doc_name.font.size = Pt(14)

    doc.add_paragraph()

    score_val = assessment.adjusted_score or assessment.clyira_score or 0.0
    score_para = doc.add_paragraph()
    score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = score_para.add_run(f"Clyira Score: {score_val:.1f} / 100")
    sr.bold = True
    sr.font.size = Pt(16)
    r, g, b = hex_to_rgb(_score_color(score_val))
    sr.font.color.rgb = RGBColor(r, g, b)

    band = assessment.score_band or "—"
    band_para = doc.add_paragraph()
    band_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = band_para.add_run(f"Band: {band}")
    br.font.size = Pt(13)
    br.bold = True

    doc.add_paragraph()

    if assessment.data_integrity_hold:
        hold_para = doc.add_paragraph()
        hold_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hold_para.add_run("⚠ DATA INTEGRITY HOLD")
        hr.bold = True
        hr.font.size = Pt(12)
        hr.font.color.rgb = RGBColor(220, 38, 38)

    doc.add_page_break()

    # ── Document Metadata ─────────────────────────────────────────────────────
    add_heading("Document Information", 1)
    for k, v in [
        ("Document", document.title),
        ("Category", document.document_category or "—"),
        ("Document No.", document.document_number or "—"),
        ("Version", document.version or "—"),
        ("Department", document.department_owner or "—"),
        ("DTAP", assessment.dtap_id or "—"),
        ("Assessment Date", assessment.created_at.strftime("%Y-%m-%d %H:%M UTC") if assessment.created_at else "—"),
        ("Levels Run", ", ".join(assessment.levels_run or [])),
        ("Processing Time", f"{assessment.processing_time_seconds:.1f}s" if assessment.processing_time_seconds else "—"),
    ]:
        p = doc.add_paragraph(style="Normal")
        add_kv(p, k, str(v))

    doc.add_paragraph()

    # ── Executive Summary ─────────────────────────────────────────────────────
    add_heading("Executive Summary", 1)

    summary_table = doc.add_table(rows=1, cols=5)
    summary_table.style = "Table Grid"
    hdr = summary_table.rows[0].cells
    for i, label in enumerate(["Critical", "High", "Medium", "Low", "Info"]):
        hdr[i].text = label
        for para in hdr[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(9)

    row = summary_table.add_row().cells
    for i, count in enumerate([
        assessment.findings_critical, assessment.findings_high,
        assessment.findings_medium, assessment.findings_low, assessment.findings_info,
    ]):
        row[i].text = str(count)
        for para in row[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(10)
                run.bold = True

    doc.add_paragraph()

    if assessment.suspended_reason:
        hold_p = doc.add_paragraph()
        r = hold_p.add_run(f"⚠ {assessment.suspended_reason}")
        r.bold = True
        r.font.color.rgb = RGBColor(220, 38, 38)
        r.font.size = Pt(10)

    doc.add_paragraph()

    # ── Findings Detail ───────────────────────────────────────────────────────
    add_heading("Assessment Findings", 1)

    sorted_findings = sorted(findings, key=lambda f: (_SEV_ORDER.get(f.severity, 5), f.level))
    open_count = sum(1 for f in findings if f.status == "open")
    resolved_count = sum(1 for f in findings if f.status == "resolved")

    stats_p = doc.add_paragraph()
    stats_p.add_run(f"{len(findings)} total findings — {open_count} open, {resolved_count} resolved")
    stats_p.runs[0].font.size = Pt(10)
    stats_p.runs[0].italic = True

    doc.add_paragraph()

    for finding in sorted_findings:
        sev = finding.severity
        sev_c = hex_to_rgb(_sev_color(sev))
        level_name = _LEVEL_NAMES.get(finding.level, finding.level)

        # Finding header
        fh = doc.add_paragraph()
        fh_run = fh.add_run(f"[{finding.level} · {level_name}] {sev.upper()} — {finding.title}")
        fh_run.bold = True
        fh_run.font.size = Pt(11)
        fh_run.font.color.rgb = RGBColor(*sev_c)

        # Status badge
        status_p = doc.add_paragraph()
        st_run = status_p.add_run(f"Status: {finding.status.upper()}")
        st_run.font.size = Pt(9)
        st_run.italic = True

        # Description
        dp = doc.add_paragraph()
        dp.add_run("Description: ").bold = True
        dp.add_run(finding.description or "").font.size = Pt(10)

        # Evidence
        if finding.evidence:
            ep = doc.add_paragraph()
            ep.add_run("Evidence: ").bold = True
            ep.add_run(f'"{finding.evidence}"').font.size = Pt(9)

        # Citation + Location row
        if finding.regulatory_citation or finding.location:
            cp = doc.add_paragraph()
            if finding.location:
                cp.add_run("Location: ").bold = True
                cp.add_run(f"{finding.location}   ").font.size = Pt(9)
            if finding.regulatory_citation:
                cp.add_run("Citation: ").bold = True
                cp.add_run(finding.regulatory_citation).font.size = Pt(9)

        # Enforcement context
        if finding.enforcement_match and finding.enforcement_context:
            efp = doc.add_paragraph()
            er = efp.add_run(f"⚡ Enforcement: {finding.enforcement_context}")
            er.font.size = Pt(9)
            er.font.color.rgb = RGBColor(185, 28, 28)

        # Remediation
        if finding.suggestion_draft:
            rp = doc.add_paragraph()
            rp.add_run("Suggested Remediation: ").bold = True
            rr = rp.add_run(finding.suggestion_draft)
            rr.font.size = Pt(10)
            rr.font.color.rgb = RGBColor(30, 64, 175)

        if finding.response_text:
            resp_p = doc.add_paragraph()
            resp_p.add_run("Response: ").bold = True
            resp_p.add_run(finding.response_text).font.size = Pt(9)

        doc.add_paragraph()  # spacer between findings

    # ── Footer ────────────────────────────────────────────────────────────────
    doc.add_page_break()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        f"Generated by Clyira Quality Intelligence Platform · {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}\n"
        "This report is generated for internal quality review purposes only."
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(107, 114, 128)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _build_redlined_docx(document_bytes: bytes, findings: list) -> bytes:
    """Return original DOCX with every suggestion_draft inserted as a w:ins tracked change."""
    import io
    from docx import Document as DocxDoc
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from datetime import datetime, timezone

    doc = DocxDoc(io.BytesIO(document_bytes))
    body = doc.element.body
    date_str = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

    # Start IDs high to avoid collision with any existing revision IDs in the document
    _id = [1000]

    SEV_COLOR = {"critical": "DC2626", "high": "EA580C", "medium": "D97706", "low": "2563EB"}

    def _next_id() -> str:
        v = str(_id[0]); _id[0] += 1; return v

    def _ins(parent) -> OxmlElement:
        el = OxmlElement("w:ins")
        el.set(qn("w:id"), _next_id())
        el.set(qn("w:author"), "Clyira AI")
        el.set(qn("w:date"), date_str)
        parent.append(el)
        return el

    def _run_with_text(text: str, bold: bool, color: str) -> OxmlElement:
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        if bold:
            rPr.append(OxmlElement("w:b"))
        clr = OxmlElement("w:color"); clr.set(qn("w:val"), color); rPr.append(clr)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text
        r.append(t)
        return r

    def _make_suggestion_paragraph(finding) -> OxmlElement:
        color = SEV_COLOR.get(finding.severity, "374151")
        p = OxmlElement("w:p")

        # Mark paragraph mark itself as inserted (required for full OOXML compliance)
        pPr = OxmlElement("w:pPr")
        rPr_pPr = OxmlElement("w:rPr")
        ins_pm = OxmlElement("w:ins")
        ins_pm.set(qn("w:id"), _next_id()); ins_pm.set(qn("w:author"), "Clyira AI"); ins_pm.set(qn("w:date"), date_str)
        rPr_pPr.append(ins_pm); pPr.append(rPr_pPr); p.append(pPr)

        # Label run: "[Clyira SEVERITY — title]: "
        ins_label = _ins(p)
        ins_label.append(_run_with_text(f"[Clyira {finding.severity.upper()} — {finding.title}]: ", True, color))

        # Suggestion text
        ins_text = _ins(p)
        ins_text.append(_run_with_text(finding.suggestion_draft or "", False, color))

        return p

    actionable = [
        f for f in findings
        if f.suggestion_draft and f.suggestion_draft.strip() and f.severity not in ("info",)
    ]

    all_paras = doc.paragraphs

    for finding in actionable:
        # Find the best matching paragraph via evidence or location text
        target = None
        needles = []
        if finding.evidence:
            needles.append(finding.evidence[:80].lower())
        if finding.location:
            needles.append(finding.location[:60].lower())

        for para in all_paras:
            pt = para.text.lower()
            if pt and any(n[:30] in pt for n in needles if n):
                target = para._element
                break

        suggestion = _make_suggestion_paragraph(finding)
        if target is not None:
            target.addnext(suggestion)
        else:
            body.append(suggestion)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@router.get("/{assessment_id}/export/redlined")
async def export_redlined_docx(
    assessment_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Download the original DOCX with all finding suggestions inserted as tracked changes."""
    import os

    result = await db.execute(
        select(Assessment).where(
            Assessment.id == assessment_id,
            Assessment.company_id == current_user.company_id,
        )
    )
    assessment = result.scalar_one_or_none()
    if not assessment:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Assessment not found")
    if assessment.status != "completed":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Assessment not completed yet")

    doc_result = await db.execute(select(Document).where(Document.id == assessment.document_id))
    document = doc_result.scalar_one_or_none()
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    if (document.file_type or "").lower() not in ("docx", "doc"):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Redlined export is only available for DOCX files")

    # Fetch raw file bytes from whichever storage backend is active
    try:
        if settings.SUPABASE_URL and settings.SUPABASE_SERVICE_KEY and document.file_path:
            from supabase import create_client
            import asyncio
            client = create_client(settings.SUPABASE_URL, settings.SUPABASE_SERVICE_KEY)
            file_bytes = await asyncio.get_event_loop().run_in_executor(
                None,
                lambda: client.storage.from_(settings.SUPABASE_STORAGE_BUCKET).download(document.file_path),
            )
        elif document.file_path and os.path.exists(document.file_path):
            with open(document.file_path, "rb") as fh:
                file_bytes = fh.read()
        else:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Original file not found in storage")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Storage retrieval failed: {exc}")

    findings_result = await db.execute(
        select(Finding).where(Finding.assessment_id == assessment_id).order_by(Finding.severity)
    )
    findings = findings_result.scalars().all()

    try:
        redlined = _build_redlined_docx(file_bytes, findings)
    except Exception as exc:
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Redlined generation failed: {exc}")

    safe = "".join(c for c in document.title if c.isalnum() or c in " _-")[:50]
    filename = f"Clyira_Redlined_{safe}_{assessment_id[:8]}.docx"

    return StreamingResponse(
        io.BytesIO(redlined),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{assessment_id}/export/docx")
async def export_assessment_docx(
    assessment_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Download assessment report as a formatted DOCX file."""
    result = await db.execute(
        select(Assessment).where(
            Assessment.id == assessment_id,
            Assessment.company_id == current_user.company_id,
        )
    )
    assessment = result.scalar_one_or_none()
    if not assessment:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Assessment not found")
    if assessment.status != "completed":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Assessment not completed yet")

    doc_result = await db.execute(select(Document).where(Document.id == assessment.document_id))
    document = doc_result.scalar_one_or_none()
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Document not found")

    findings_result = await db.execute(
        select(Finding)
        .where(Finding.assessment_id == assessment_id)
        .order_by(Finding.severity)
    )
    findings = findings_result.scalars().all()

    try:
        docx_bytes = _build_docx(assessment, document, findings)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Report generation failed: {e}",
        )

    safe_title = "".join(c for c in document.title if c.isalnum() or c in " _-")[:50]
    filename = f"Clyira_{safe_title}_{assessment_id[:8]}.docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{assessment_id}/export/csv")
async def export_assessment_csv(
    assessment_id: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Download findings as a CSV — suitable for QMS/LIMS import or Excel analysis."""
    import csv

    result = await db.execute(
        select(Assessment).where(
            Assessment.id == assessment_id,
            Assessment.company_id == current_user.company_id,
        )
    )
    assessment = result.scalar_one_or_none()
    if not assessment:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Assessment not found")
    if assessment.status != "completed":
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Assessment not completed yet")

    doc_result = await db.execute(select(Document).where(Document.id == assessment.document_id))
    document = doc_result.scalar_one_or_none()

    findings_result = await db.execute(
        select(Finding)
        .where(Finding.assessment_id == assessment_id)
        .order_by(Finding.severity)
    )
    findings = findings_result.scalars().all()

    buf = io.StringIO()
    writer = csv.writer(buf)

    writer.writerow([
        "finding_id", "assessment_id", "document_title", "document_number",
        "level", "level_name", "severity", "category", "title", "description",
        "evidence", "location", "regulatory_citation", "citation_type", "agency",
        "enforcement_match", "severity_elevated", "status",
        "suggestion_draft", "next_step_text", "remediation_priority",
        "confidence_score", "validated",
    ])

    doc_title = document.title if document else ""
    doc_number = document.document_number if document else ""

    for f in findings:
        writer.writerow([
            f.id, assessment_id, doc_title, doc_number,
            f.level, _LEVEL_NAMES.get(f.level or "", f.level or ""),
            f.severity, f.category or "", f.title,
            (f.description or "").replace("\n", " "),
            (f.evidence or "").replace("\n", " "),
            f.location or "", f.regulatory_citation or "",
            f.citation_type or "", f.agency or "",
            "Yes" if f.enforcement_match else "No",
            "Yes" if f.severity_elevated else "No",
            f.status,
            (f.suggestion_draft or "").replace("\n", " "),
            (f.next_step_text or "").replace("\n", " "),
            f.remediation_priority or "",
            f.confidence_score or "",
            "Yes" if f.validated else "No",
        ])

    safe_title = "".join(c for c in (document.title if document else "report") if c.isalnum() or c in " _-")[:50]
    filename = f"Clyira_Findings_{safe_title}_{assessment_id[:8]}.csv"
    csv_bytes = buf.getvalue().encode("utf-8-sig")  # utf-8-sig for Excel BOM

    return StreamingResponse(
        io.BytesIO(csv_bytes),
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
