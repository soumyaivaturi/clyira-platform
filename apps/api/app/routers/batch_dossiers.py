"""
Batch & Lot Record Review — API Router
Provides CRUD for BatchDossiers, document linking, readiness computation,
disposition decisions, and evidence completeness checks.
"""
import logging
from datetime import datetime, timezone

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.database import get_db
from app.core.dependencies import get_current_user
from app.models.batch_dossier import (
    BatchDossier, BatchDossierDocument, EvidencePackageTemplate, FeedbackCorrection
)
from app.models.assessment import Assessment, Finding
from app.models.document import Document
from app.models.base import generate_uuid
from app.models.user import User
from app.schemas.batch_dossier import (
    BatchDossierCreate, BatchDossierUpdate,
    BatchDossierDocumentAdd, DispositionDecisionCreate,
    FeedbackCorrectionCreate,
)
from app.services.batch_disposition_service import BatchDispositionService
from app.services.evidence_completeness_service import EvidenceCompletenessService

logger = logging.getLogger(__name__)
router = APIRouter()


# ── Document scanning (field extraction for form pre-fill) ───────────────────

_BPR_FIELDS = ["lot_number", "product_name", "product_code", "dosage_form",
               "batch_size", "manufacturing_site", "manufacturing_date", "target_release_date"]

_BPR_LLM_SYSTEM = (
    "You are a pharmaceutical document parser specialising in batch records from CDMOs and manufacturers "
    "(Catalent, Lonza, Patheon, etc.). Extract header fields from the text. "
    "Field mapping hints: 'Batch Record' or 'Record Number' → lot_number; "
    "'Customer Protocol' or 'Item Description' → product_name; "
    "'Item Number' or 'Project Number' → product_code; "
    "'Catalent Site' or 'Facility' → manufacturing_site; "
    "'Planned Output Quantity' → batch_size; "
    "'Expiry Date' → target_release_date; 'Effective Date' or 'Start Date' → manufacturing_date. "
    "Return ONLY a JSON object with these exact keys (use null if not found): "
    "lot_number, product_name, product_code, dosage_form, batch_size, "
    "manufacturing_site, manufacturing_date (YYYY-MM-DD), target_release_date (YYYY-MM-DD). "
    "No explanation, no markdown — just the JSON object."
)


_BPR_VISION_PROMPT = (
    "This is a pharmaceutical batch record (MBR/BPR) — possibly from Catalent or another CDMO. "
    "Read the LABELED TEXT FIELDS only. Do NOT read barcodes, QR codes, GS1 codes, or any long strings of digits without dashes. "
    "Pharma lot/batch numbers look like '37040001-XBR-8-1' or 'BPR-2024-001A' — alphanumeric with dashes. "
    "Field mapping hints: 'Batch Record:'/'Record Number:' → LOT_NUMBER; "
    "'Item Description'/'Customer Protocol'/'Customer:' → PRODUCT_NAME; "
    "'Item Number'/'Project Number:' → PRODUCT_CODE; "
    "'Catalent Site:'/'Facility:'/'Site:' → MANUFACTURING_SITE; "
    "'Planned Output Quantity'/'Batch Size:' → BATCH_SIZE; "
    "'Expiry Date:' → TARGET_RELEASE_DATE; 'Effective Date:'/'Start Date:' → MANUFACTURING_DATE. "
    "Reply ONLY with lines in this exact format (omit any field you cannot clearly read):\n"
    "LOT_NUMBER: <value>\n"
    "PRODUCT_NAME: <value>\n"
    "PRODUCT_CODE: <value>\n"
    "DOSAGE_FORM: <value>\n"
    "BATCH_SIZE: <value>\n"
    "MANUFACTURING_SITE: <value>\n"
    "MANUFACTURING_DATE: <YYYY-MM-DD>\n"
    "TARGET_RELEASE_DATE: <YYYY-MM-DD>\n"
    "No extra text, no JSON, no markdown."
)


async def _vision_scan_bpr_fields(pdf_bytes: bytes) -> tuple[dict, str | None]:
    """Extract BPR fields from a scanned/screenshot PDF using Gemini Vision.
    Returns (fields_dict, error_string). error_string is None on success.
    """
    import base64, httpx, re
    from app.core.config import settings

    try:
        import fitz
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        # Scan only page 1 — BPR header fields are always on the cover page.
        # Keeping to 1 page at 150 DPI keeps image size under Gemini limits and
        # avoids OOM on Render's free tier (512 MB RAM).
        image_parts = []
        page = doc[0]
        pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), colorspace=fitz.csGRAY)  # 144 DPI greyscale
        img_b64 = base64.b64encode(pix.tobytes("jpeg", jpg_quality=75)).decode()
        image_parts.append({"inline_data": {"mime_type": "image/jpeg", "data": img_b64}})
        doc.close()
        logger.info(f"Vision: rendered page 1 for Gemini, b64 size={len(img_b64)}")
    except Exception as e:
        msg = f"PDF→image failed: {e}"
        logger.warning(f"Vision scan: {msg}")
        return {}, msg

    url = (
        f"https://generativelanguage.googleapis.com/v1beta/models/"
        f"{settings.GEMINI_MODEL}:generateContent?key={settings.GEMINI_API_KEY}"
    )
    payload = {
        "contents": [{"role": "user", "parts": [{"text": _BPR_VISION_PROMPT}] + image_parts}],
        "generationConfig": {"temperature": 0.1, "maxOutputTokens": 1024},
    }
    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            resp = await client.post(url, json=payload)
            if resp.status_code != 200:
                msg = f"Gemini HTTP {resp.status_code}: {resp.text[:200]}"
                logger.warning(f"Vision scan: {msg}")
                return {}, msg
            raw = resp.json()["candidates"][0]["content"]["parts"][0]["text"]
            logger.info(f"Gemini raw response (first 300): {raw[:300]}")
            # Parse KEY: VALUE lines — one per field, ASCII only, no JSON
            _KEY_MAP = {
                "LOT_NUMBER": "lot_number",
                "PRODUCT_NAME": "product_name",
                "PRODUCT_CODE": "product_code",
                "DOSAGE_FORM": "dosage_form",
                "BATCH_SIZE": "batch_size",
                "MANUFACTURING_SITE": "manufacturing_site",
                "MANUFACTURING_DATE": "manufacturing_date",
                "TARGET_RELEASE_DATE": "target_release_date",
            }
            _SKIP_VALUES = {"<value>", "<yyyy-mm-dd>", "null", "n/a", "none", "unknown", ""}
            result = {}
            for line in raw.splitlines():
                if ":" not in line:
                    continue
                key, _, val = line.partition(":")
                key = key.strip().upper()
                val = val.strip()
                if key not in _KEY_MAP or val.lower() in _SKIP_VALUES:
                    continue
                # Reject lot numbers that are pure digits >10 chars — almost certainly a misread barcode
                if key == "LOT_NUMBER" and re.fullmatch(r'\d{10,}', val):
                    logger.info(f"Vision: rejected probable barcode as lot_number: {val[:20]}")
                    continue
                result[_KEY_MAP[key]] = val
            logger.info(f"Gemini Vision extracted: {result}")
            return result, None
    except Exception as e:
        msg = f"Gemini call failed: {e}"
        logger.warning(f"Vision scan: {msg}")
        return {}, msg


async def _llm_extract_bpr_fields(text: str) -> dict:
    """LLM fallback for BPR field extraction when regex finds nothing (text already extracted)."""
    import json
    from app.core.config import settings
    from app.engines.llm_engine import _call_groq, _call_gemini

    snippet = text[:4000]
    try:
        if settings.GROQ_API_KEY:
            raw = await _call_groq(_BPR_LLM_SYSTEM, f"Extract fields from this batch record:\n\n{snippet}")
        else:
            raw = await _call_gemini(_BPR_LLM_SYSTEM, f"Extract fields from this batch record:\n\n{snippet}")
        raw = raw.strip().removeprefix("```json").removeprefix("```").removesuffix("```").strip()
        data = json.loads(raw)
        return {k: str(v).strip() if v not in (None, "", "null") else None for k, v in data.items()
                if k in _BPR_FIELDS}
    except Exception as e:
        logger.warning(f"LLM BPR extraction failed: {e}")
        return {}


@router.post("/scan-document")
async def scan_document_for_fields(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
):
    """
    Extract BPR header fields from an uploaded file without creating a document record.
    Extraction chain: IDPEngine (auto-routed) → regex patterns → LLM fallback.
    The file is NOT persisted — call the normal document upload endpoint separately.
    """
    from app.services.bpr_extraction_service import BPRExtractionService

    content = await file.read()
    filename = file.filename or "upload"
    file_type = filename.rsplit(".", 1)[-1].lower() if "." in filename else "unknown"

    import io

    # Scan uses fast extractors only (no Tesseract — too slow for form pre-fill).
    # IDPEngine with OCR runs only in the full assessment pipeline (background task).
    extracted_text = ""

    # Step 1a: DOCX via python-docx
    if file_type in ("docx", "doc"):
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(content))
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            extracted_text = "\n".join(parts)
            logger.info(f"python-docx: {len(extracted_text)} chars from {filename!r}")
        except Exception as e:
            logger.warning(f"python-docx failed: {e}")

    # Step 1b: pdfplumber for native/digital PDFs — first 3 pages only (header is always page 1)
    if not extracted_text.strip() and file_type == "pdf":
        try:
            import pdfplumber
            pages = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages[:3]:  # cap at 3 pages — enough for any header
                    t = page.extract_text() or ""
                    for table in page.extract_tables():
                        rows = [" | ".join(str(c or "").strip() for c in row) for row in table if any(c for c in row)]
                        t += "\n" + "\n".join(rows)
                    if t.strip():
                        pages.append(t)
            extracted_text = "\n".join(pages)
            logger.info(f"pdfplumber: {len(extracted_text)} chars from {filename!r}")
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")

    # Step 1d: Gemini Vision — for scanned/screenshot PDFs where no text could be extracted.
    # Returns structured fields directly (skips regex + LLM text path).
    vision_error = None
    if not extracted_text.strip() and file_type == "pdf":
        from app.core.config import settings
        if not settings.GEMINI_API_KEY:
            vision_error = "GEMINI_API_KEY not set"
            logger.warning("Vision scan skipped: GEMINI_API_KEY not set on this deployment")
        else:
            vision_fields, vision_error = await _vision_scan_bpr_fields(content)
            if vision_fields:
                fields_dict = {k: vision_fields.get(k) for k in _BPR_FIELDS}
                confidence_dict = {k: (0.82 if vision_fields.get(k) else None) for k in _BPR_FIELDS}
                fields_found = sum(1 for v in fields_dict.values() if v)
                logger.info(f"BPR scan via vision: file={filename!r} fields_found={fields_found}")
                return {
                    "fields": fields_dict,
                    "confidence": confidence_dict,
                    "filename": filename,
                    "file_type": file_type,
                    "text_length": 0,
                    "extraction_method": "vision",
                    "fields_found": fields_found,
                }

    # Step 2: Regex extraction
    fields = BPRExtractionService().extract(extracted_text)
    fields_dict = {
        "lot_number": fields.lot_number.value if fields.lot_number else None,
        "product_name": fields.product_name.value if fields.product_name else None,
        "product_code": fields.product_code.value if fields.product_code else None,
        "dosage_form": fields.dosage_form.value if fields.dosage_form else None,
        "batch_size": fields.batch_size.value if fields.batch_size else None,
        "manufacturing_site": fields.manufacturing_site.value if fields.manufacturing_site else None,
        "manufacturing_date": fields.manufacturing_date.value if fields.manufacturing_date else None,
        "target_release_date": fields.target_release_date.value if fields.target_release_date else None,
    }
    confidence_dict = {
        "lot_number": fields.lot_number.confidence if fields.lot_number else None,
        "product_name": fields.product_name.confidence if fields.product_name else None,
        "product_code": fields.product_code.confidence if fields.product_code else None,
        "dosage_form": fields.dosage_form.confidence if fields.dosage_form else None,
        "batch_size": fields.batch_size.confidence if fields.batch_size else None,
        "manufacturing_site": fields.manufacturing_site.confidence if fields.manufacturing_site else None,
        "manufacturing_date": fields.manufacturing_date.confidence if fields.manufacturing_date else None,
        "target_release_date": fields.target_release_date.confidence if fields.target_release_date else None,
    }

    # Step 3: LLM fallback — fires when regex found nothing but text exists (≥50 chars)
    regex_hit_count = sum(1 for v in fields_dict.values() if v)
    extraction_method = "regex"
    if regex_hit_count == 0 and len(extracted_text) >= 50:
        llm_fields = await _llm_extract_bpr_fields(extracted_text)
        if llm_fields:
            for k, v in llm_fields.items():
                if v and fields_dict.get(k) is None:
                    fields_dict[k] = v
                    confidence_dict[k] = 0.75  # LLM extractions default to 0.75 confidence
            extraction_method = "llm"

    fields_found = sum(1 for v in fields_dict.values() if v)
    logger.info(
        f"BPR scan complete: file={filename!r} text_len={len(extracted_text)} "
        f"method={extraction_method} fields_found={fields_found} vision_error={vision_error!r}"
    )
    return {
        "fields": fields_dict,
        "confidence": confidence_dict,
        "filename": filename,
        "file_type": file_type,
        "text_length": len(extracted_text),
        "extraction_method": extraction_method,
        "fields_found": fields_found,
        "debug": vision_error,  # shows why vision was skipped/failed
    }


# ── Serialisers ──────────────────────────────────────────────────────────────

def _dossier_out(d: BatchDossier, documents: list = None, readiness: dict = None) -> dict:
    return {
        "id": d.id,
        "company_id": d.company_id,
        "created_by": d.created_by,
        "lot_number": d.lot_number,
        "product_name": d.product_name,
        "product_code": d.product_code,
        "dosage_form": d.dosage_form,
        "batch_size": d.batch_size,
        "manufacturing_site": d.manufacturing_site,
        "manufacturing_date": d.manufacturing_date,
        "target_release_date": d.target_release_date,
        # Layer 0
        "record_family": d.record_family,
        "product_type": d.product_type,
        "is_sterile": d.is_sterile,
        "manufacturing_context": d.manufacturing_context,
        "batch_purpose": d.batch_purpose,
        "target_markets": d.target_markets or [],
        # Status
        "status": d.status,
        "readiness_status": d.readiness_status,
        "readiness_score": d.readiness_score,
        "readiness_band": d.readiness_band,
        # Disposition
        "disposition_decision": d.disposition_decision,
        "disposition_rationale": d.disposition_rationale,
        "disposition_divergence": d.disposition_divergence,
        "conditional_release_conditions": d.conditional_release_conditions,
        # Gates
        "gates": {
            "evidence_complete": d.gate_evidence_complete,
            "open_deviations": d.gate_open_deviations,
            "open_capas": d.gate_open_capas,
            "qc_complete": d.gate_qc_complete,
            "data_integrity_ok": not d.gate_data_integrity,
            "all_findings_addressed": d.gate_all_findings_addressed,
            "gray_findings_resolved": d.gate_gray_findings_resolved,
        },
        # Review
        "shadow_mode": d.shadow_mode,
        "review_stage": d.review_stage,
        "released_by": d.released_by,
        "released_at": d.released_at,
        "documents": documents or [],
        "readiness_detail": readiness or {},
        "created_at": str(d.created_at),
        "updated_at": str(d.updated_at),
    }


def _dossier_doc_out(dd: BatchDossierDocument) -> dict:
    return {
        "id": dd.id,
        "dossier_id": dd.dossier_id,
        "document_id": dd.document_id,
        "role": dd.role,
        "sequence_order": dd.sequence_order,
        "notes": dd.notes,
        "added_by": dd.added_by,
        "added_at": str(dd.created_at),
    }


def _finding_out(f: Finding) -> dict:
    return {
        "id": f.id,
        "level": f.level,
        "level_name": f.level_name,
        "severity": f.severity,
        "category": f.category,
        "title": f.title,
        "description": f.description,
        "evidence": f.evidence,
        "location": f.location,
        "regulatory_citation": f.regulatory_citation,
        "enforcement_match": f.enforcement_match,
        "suggestion_draft": f.suggestion_draft,
        "status": f.status,
        "confidence_score": f.confidence_score,
        "verification_state": getattr(f, "verification_state", None),
        "field_criticality": getattr(f, "field_criticality", None),
        "source_page": getattr(f, "source_page", None),
        "human_verification_required": getattr(f, "human_verification_required", False),
        "extraction_confidence": getattr(f, "extraction_confidence", None),
        "explanation_trace": getattr(f, "explanation_trace", None),
    }


# ── Endpoints ─────────────────────────────────────────────────────────────────

@router.get("")
async def list_dossiers(
    status: str | None = None,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """List all batch dossiers for the current user's company."""
    q = select(BatchDossier).where(BatchDossier.company_id == current_user.company_id)
    if status:
        q = q.where(BatchDossier.status == status)
    q = q.order_by(BatchDossier.created_at.desc())
    result = await db.execute(q)
    dossiers = result.scalars().all()

    output = []
    for d in dossiers:
        docs_result = await db.execute(
            select(BatchDossierDocument).where(BatchDossierDocument.dossier_id == d.id)
        )
        docs = docs_result.scalars().all()
        output.append(_dossier_out(d, [_dossier_doc_out(dd) for dd in docs]))
    return output


@router.post("", status_code=status.HTTP_201_CREATED)
async def create_dossier(
    body: BatchDossierCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Create a new batch dossier with Layer 0 classification."""
    dossier = BatchDossier(
        id=generate_uuid(),
        company_id=current_user.company_id,
        created_by=current_user.id,
        lot_number=body.lot_number,
        product_name=body.product_name,
        product_code=body.product_code,
        dosage_form=body.dosage_form,
        batch_size=body.batch_size,
        manufacturing_site=body.manufacturing_site,
        manufacturing_date=body.manufacturing_date,
        target_release_date=body.target_release_date,
        record_family=body.record_family,
        product_type=body.product_type,
        is_sterile=body.is_sterile,
        manufacturing_context=body.manufacturing_context,
        batch_purpose=body.batch_purpose,
        target_markets=body.target_markets,
        shadow_mode=body.shadow_mode,
        status="draft",
    )
    db.add(dossier)
    await db.commit()
    await db.refresh(dossier)
    return _dossier_out(dossier)


@router.get("/{dossier_id}")
async def get_dossier(
    dossier_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Get a dossier with full detail: documents, findings, readiness."""
    dossier = await db.get(BatchDossier, dossier_id)
    if not dossier or dossier.company_id != current_user.company_id:
        raise HTTPException(status_code=404, detail="Dossier not found")

    docs_result = await db.execute(
        select(BatchDossierDocument).where(BatchDossierDocument.dossier_id == dossier_id)
    )
    dossier_docs = docs_result.scalars().all()

    # Enrich each dossier document with assessment summary and findings
    enriched_docs = []
    for dd in dossier_docs:
        dd_out = _dossier_doc_out(dd)

        # Load document title
        doc = await db.get(Document, dd.document_id)
        if doc:
            dd_out["document_title"] = doc.title
            dd_out["document_category"] = doc.document_category
            dd_out["document_status"] = doc.status

        # Load latest completed assessment
        assessment_result = await db.execute(
            select(Assessment)
            .where(Assessment.document_id == dd.document_id)
            .where(Assessment.status == "completed")
            .order_by(Assessment.created_at.desc())
            .limit(1)
        )
        assessment = assessment_result.scalar_one_or_none()
        if assessment:
            dd_out["assessment"] = {
                "id": assessment.id,
                "clyira_score": assessment.clyira_score,
                "score_band": assessment.score_band,
                "findings_critical": assessment.findings_critical,
                "findings_high": assessment.findings_high,
                "findings_medium": assessment.findings_medium,
                "findings_low": assessment.findings_low,
                "completed_at": assessment.completed_at,
            }

            # Load findings for this document
            findings_result = await db.execute(
                select(Finding).where(Finding.assessment_id == assessment.id)
                .order_by(Finding.severity)
            )
            findings = findings_result.scalars().all()
            dd_out["findings"] = [_finding_out(f) for f in findings]
        else:
            dd_out["assessment"] = None
            dd_out["findings"] = []

        enriched_docs.append(dd_out)

    # Evidence completeness
    svc = EvidenceCompletenessService()
    ev_check = svc.check(dossier, dossier_docs)

    return _dossier_out(dossier, enriched_docs, ev_check)


@router.patch("/{dossier_id}")
async def update_dossier(
    dossier_id: str,
    body: BatchDossierUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Update dossier metadata or Layer 0 classification."""
    dossier = await db.get(BatchDossier, dossier_id)
    if not dossier or dossier.company_id != current_user.company_id:
        raise HTTPException(status_code=404, detail="Dossier not found")

    for field, value in body.model_dump(exclude_unset=True).items():
        setattr(dossier, field, value)

    await db.commit()
    await db.refresh(dossier)
    return _dossier_out(dossier)


@router.delete("/{dossier_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_dossier(
    dossier_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Delete a draft dossier. Only allowed when status=draft."""
    dossier = await db.get(BatchDossier, dossier_id)
    if not dossier or dossier.company_id != current_user.company_id:
        raise HTTPException(status_code=404, detail="Dossier not found")
    if dossier.status != "draft":
        raise HTTPException(status_code=400, detail="Only draft dossiers can be deleted")
    await db.delete(dossier)
    await db.commit()


@router.post("/{dossier_id}/documents")
async def add_document(
    dossier_id: str,
    body: BatchDossierDocumentAdd,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Add a document to a dossier with a role."""
    dossier = await db.get(BatchDossier, dossier_id)
    if not dossier or dossier.company_id != current_user.company_id:
        raise HTTPException(status_code=404, detail="Dossier not found")

    # Verify document belongs to same company
    doc = await db.get(Document, body.document_id)
    if not doc or doc.company_id != current_user.company_id:
        raise HTTPException(status_code=404, detail="Document not found")

    # Check for duplicate
    existing = await db.execute(
        select(BatchDossierDocument)
        .where(BatchDossierDocument.dossier_id == dossier_id)
        .where(BatchDossierDocument.document_id == body.document_id)
    )
    if existing.scalar_one_or_none():
        raise HTTPException(status_code=400, detail="Document already added to this dossier")

    dd = BatchDossierDocument(
        id=generate_uuid(),
        dossier_id=dossier_id,
        document_id=body.document_id,
        role=body.role,
        sequence_order=body.sequence_order,
        notes=body.notes,
        added_by=current_user.id,
    )
    db.add(dd)

    # Advance status to under_review if still draft
    if dossier.status == "draft":
        dossier.status = "under_review"

    await db.commit()
    await db.refresh(dd)

    # Recompute readiness with new document set
    try:
        svc = BatchDispositionService(db)
        await svc.compute_readiness(dossier_id)
    except Exception as e:
        logger.warning(f"Readiness recompute after add_document failed: {e}")

    return _dossier_doc_out(dd)


@router.delete("/{dossier_id}/documents/{dossier_doc_id}", status_code=status.HTTP_204_NO_CONTENT)
async def remove_document(
    dossier_id: str,
    dossier_doc_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Remove a document from a dossier."""
    dossier = await db.get(BatchDossier, dossier_id)
    if not dossier or dossier.company_id != current_user.company_id:
        raise HTTPException(status_code=404, detail="Dossier not found")

    dd = await db.get(BatchDossierDocument, dossier_doc_id)
    if not dd or dd.dossier_id != dossier_id:
        raise HTTPException(status_code=404, detail="Document link not found")

    await db.delete(dd)
    await db.commit()


@router.post("/{dossier_id}/assess-readiness")
async def assess_readiness(
    dossier_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Trigger readiness computation for a dossier.
    This runs the BatchDispositionService and updates all gate flags and readiness status.
    """
    dossier = await db.get(BatchDossier, dossier_id)
    if not dossier or dossier.company_id != current_user.company_id:
        raise HTTPException(status_code=404, detail="Dossier not found")

    svc = BatchDispositionService(db)
    result = await svc.compute_readiness(dossier_id)
    return result


@router.post("/{dossier_id}/disposition")
async def record_disposition(
    dossier_id: str,
    body: DispositionDecisionCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Record the QA Approver's final disposition decision.
    Validates that rationale is provided and flags any divergence from readiness status.
    """
    dossier = await db.get(BatchDossier, dossier_id)
    if not dossier or dossier.company_id != current_user.company_id:
        raise HTTPException(status_code=404, detail="Dossier not found")

    if body.decision not in ("release", "conditional_release", "hold", "reject"):
        raise HTTPException(status_code=400, detail="Invalid disposition decision")

    if len(body.rationale.strip()) < 20:
        raise HTTPException(status_code=400, detail="Disposition rationale must be at least 20 characters")

    svc = BatchDispositionService(db)
    result = await svc.record_disposition_decision(
        dossier_id=dossier_id,
        decision=body.decision,
        rationale=body.rationale,
        decided_by=current_user.id,
        conditional_conditions=body.conditional_conditions,
    )
    return result


@router.get("/{dossier_id}/findings")
async def list_dossier_findings(
    dossier_id: str,
    verification_state: str | None = None,
    severity: str | None = None,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    List all findings across all documents in a dossier.
    Optionally filter by verification_state (green/red/blue/gray) or severity.
    """
    dossier = await db.get(BatchDossier, dossier_id)
    if not dossier or dossier.company_id != current_user.company_id:
        raise HTTPException(status_code=404, detail="Dossier not found")

    docs_result = await db.execute(
        select(BatchDossierDocument).where(BatchDossierDocument.dossier_id == dossier_id)
    )
    dossier_docs = docs_result.scalars().all()

    all_findings = []
    for dd in dossier_docs:
        assessment_result = await db.execute(
            select(Assessment)
            .where(Assessment.document_id == dd.document_id)
            .where(Assessment.status == "completed")
            .order_by(Assessment.created_at.desc())
            .limit(1)
        )
        assessment = assessment_result.scalar_one_or_none()
        if not assessment:
            continue

        findings_q = select(Finding).where(Finding.assessment_id == assessment.id)
        if verification_state:
            findings_q = findings_q.where(Finding.verification_state == verification_state)
        if severity:
            findings_q = findings_q.where(Finding.severity == severity)

        findings_result = await db.execute(findings_q.order_by(Finding.severity))
        findings = findings_result.scalars().all()

        for f in findings:
            f_out = _finding_out(f)
            f_out["document_id"] = dd.document_id
            f_out["document_role"] = dd.role
            all_findings.append(f_out)

    # Sort: critical first, then high, medium, low, info
    severity_order = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}
    all_findings.sort(key=lambda f: severity_order.get(f["severity"], 5))

    return {
        "dossier_id": dossier_id,
        "total": len(all_findings),
        "findings": all_findings,
    }


@router.patch("/{dossier_id}/findings/{finding_id}/review")
async def review_finding(
    dossier_id: str,
    finding_id: str,
    body: dict,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Update a finding's verification state during batch record review."""
    dossier = await db.get(BatchDossier, dossier_id)
    if not dossier or dossier.company_id != current_user.company_id:
        raise HTTPException(status_code=404, detail="Dossier not found")

    finding = await db.get(Finding, finding_id)
    if not finding:
        raise HTTPException(status_code=404, detail="Finding not found")

    new_state = body.get("verification_state")
    if new_state not in ("green", "red", "gray", "blue"):
        raise HTTPException(status_code=400, detail="verification_state must be green, red, gray, or blue")

    finding.verification_state = new_state
    await db.commit()
    return {"id": finding_id, "verification_state": new_state}


@router.post("/{dossier_id}/feedback-correction")
async def submit_feedback_correction(
    dossier_id: str,
    body: FeedbackCorrectionCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Submit a reviewer correction to an AI-extracted or AI-assessed value."""
    dossier = await db.get(BatchDossier, dossier_id)
    if not dossier or dossier.company_id != current_user.company_id:
        raise HTTPException(status_code=404, detail="Dossier not found")

    correction = FeedbackCorrection(
        id=generate_uuid(),
        finding_id=body.finding_id,
        document_id=body.document_id,
        corrected_by=current_user.id,
        field_name=body.field_name,
        original_value=body.original_value,
        corrected_value=body.corrected_value,
        source_page=body.source_page,
        field_criticality=body.field_criticality,
        correction_rationale=body.correction_rationale,
    )
    db.add(correction)
    await db.commit()
    return {"id": correction.id, "status": "recorded"}


@router.get("/stats/summary")
async def dossier_stats(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Dashboard stats: counts by status and readiness for this company."""
    result = await db.execute(
        select(BatchDossier)
        .where(BatchDossier.company_id == current_user.company_id)
    )
    dossiers = result.scalars().all()

    status_counts: dict = {}
    readiness_counts: dict = {}
    for d in dossiers:
        status_counts[d.status] = status_counts.get(d.status, 0) + 1
        if d.readiness_status:
            readiness_counts[d.readiness_status] = readiness_counts.get(d.readiness_status, 0) + 1

    return {
        "total": len(dossiers),
        "by_status": status_counts,
        "by_readiness": readiness_counts,
    }


@router.post("/{dossier_id}/reopen")
async def reopen_dossier(
    dossier_id: str,
    body: dict,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Reopen a dispositioned dossier (§22.5).
    Requires a documented reason of at least 100 characters.
    """
    dossier = await db.get(BatchDossier, dossier_id)
    if not dossier or dossier.company_id != current_user.company_id:
        raise HTTPException(status_code=404, detail="Dossier not found")

    reason = (body.get("reason") or "").strip()
    if len(reason) < 100:
        raise HTTPException(
            status_code=400,
            detail="Reopen reason must be at least 100 characters to meet §22.5 documentation requirements.",
        )

    svc = BatchDispositionService(db)
    result = await svc.reopen_dossier(
        dossier_id=dossier_id,
        reason=reason,
        reopened_by=current_user.id,
    )
    if "error" in result:
        raise HTTPException(status_code=400, detail=result["error"])
    return result


@router.get("/{dossier_id}/conflicts")
async def detect_conflicts(
    dossier_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Cross-document conflict detection (§22.2).
    Returns lot number / date inconsistencies found across all documents in the dossier.
    """
    dossier = await db.get(BatchDossier, dossier_id)
    if not dossier or dossier.company_id != current_user.company_id:
        raise HTTPException(status_code=404, detail="Dossier not found")

    svc = BatchDispositionService(db)
    conflicts = await svc.detect_cross_document_conflicts(dossier_id)
    return {
        "dossier_id": dossier_id,
        "conflict_count": len(conflicts),
        "has_critical_conflicts": any(c.get("severity") == "critical" for c in conflicts),
        "conflicts": conflicts,
    }


@router.get("/{dossier_id}/report")
async def generate_review_report(
    dossier_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Generate an audit-ready review report for the dossier.
    Returns structured JSON suitable for rendering as a GMP review record.
    """
    from fastapi.responses import JSONResponse

    dossier = await db.get(BatchDossier, dossier_id)
    if not dossier or dossier.company_id != current_user.company_id:
        raise HTTPException(status_code=404, detail="Dossier not found")

    docs_result = await db.execute(
        select(BatchDossierDocument).where(BatchDossierDocument.dossier_id == dossier_id)
    )
    dossier_docs = docs_result.scalars().all()

    # Build document summaries with findings
    document_sections = []
    all_findings_flat = []
    for dd in dossier_docs:
        doc = await db.get(Document, dd.document_id)
        doc_section: dict = {
            "document_id": dd.document_id,
            "document_title": doc.title if doc else dd.document_id,
            "document_category": doc.document_category if doc else "",
            "role": dd.role,
            "assessment": None,
            "findings_summary": {},
            "open_findings": [],
        }

        assessment_result = await db.execute(
            select(Assessment)
            .where(Assessment.document_id == dd.document_id)
            .where(Assessment.status == "completed")
            .order_by(Assessment.created_at.desc())
            .limit(1)
        )
        assessment = assessment_result.scalar_one_or_none()
        if assessment:
            doc_section["assessment"] = {
                "clyira_score": assessment.clyira_score,
                "score_band": assessment.score_band,
                "data_integrity_hold": assessment.data_integrity_hold,
                "completed_at": assessment.completed_at,
            }

            findings_result = await db.execute(
                select(Finding).where(Finding.assessment_id == assessment.id)
            )
            findings = findings_result.scalars().all()

            counts: dict = {}
            open_f = []
            for f in findings:
                counts[f.severity] = counts.get(f.severity, 0) + 1
                if f.status in ("open", "disputed", "acknowledged"):
                    open_f.append({
                        "id": f.id,
                        "level": f.level,
                        "severity": f.severity,
                        "title": f.title,
                        "status": f.status,
                        "verification_state": f.verification_state,
                        "regulatory_citation": f.regulatory_citation,
                    })
                    all_findings_flat.append(f)

            doc_section["findings_summary"] = counts
            doc_section["open_findings"] = open_f

        document_sections.append(doc_section)

    svc = EvidenceCompletenessService()
    ev_check = svc.check(dossier, dossier_docs)

    now = datetime.now(timezone.utc).isoformat()
    report = {
        "report_type": "batch_dossier_review",
        "generated_at": now,
        "generated_by": current_user.id,
        "dossier": {
            "id": dossier.id,
            "lot_number": dossier.lot_number,
            "product_name": dossier.product_name,
            "product_code": dossier.product_code,
            "dosage_form": dossier.dosage_form,
            "batch_size": dossier.batch_size,
            "manufacturing_site": dossier.manufacturing_site,
            "manufacturing_date": dossier.manufacturing_date,
            "target_release_date": dossier.target_release_date,
            "is_sterile": dossier.is_sterile,
            "manufacturing_context": dossier.manufacturing_context,
            "product_type": dossier.product_type,
            "target_markets": dossier.target_markets or [],
        },
        "readiness": {
            "status": dossier.readiness_status,
            "score": dossier.readiness_score,
            "band": dossier.readiness_band,
            "gates": {
                "evidence_complete": dossier.gate_evidence_complete,
                "data_integrity_ok": not dossier.gate_data_integrity,
                "all_findings_addressed": dossier.gate_all_findings_addressed,
                "gray_findings_resolved": dossier.gate_gray_findings_resolved,
            },
        },
        "evidence_completeness": ev_check,
        "disposition": {
            "decision": dossier.disposition_decision,
            "rationale": dossier.disposition_rationale,
            "decided_by": dossier.released_by,
            "decided_at": dossier.released_at,
            "divergence_flagged": dossier.disposition_divergence,
            "conditional_conditions": dossier.conditional_release_conditions,
            "e_signature": {
                "signer_id": getattr(dossier, "disposition_signer_id", None),
                "signed_at": getattr(dossier, "disposition_signed_at", None),
                "meaning": getattr(dossier, "disposition_signature_meaning", None),
            },
        },
        "reopen_history": {
            "reopen_count": getattr(dossier, "reopen_count", 0),
            "last_reopened_by": getattr(dossier, "reopened_by", None),
            "last_reopened_at": getattr(dossier, "reopened_at", None),
            "last_reopen_reason": getattr(dossier, "reopen_reason", None),
        },
        "document_sections": document_sections,
        "finding_totals": {
            "total_open": len(all_findings_flat),
            "critical_open": sum(1 for f in all_findings_flat if f.severity == "critical"),
            "high_open": sum(1 for f in all_findings_flat if f.severity == "high"),
            "medium_open": sum(1 for f in all_findings_flat if f.severity == "medium"),
            "low_open": sum(1 for f in all_findings_flat if f.severity == "low"),
        },
    }
    return report
